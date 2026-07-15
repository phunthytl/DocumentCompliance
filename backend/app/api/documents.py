import os
import shutil
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.models.document import Document
from app.models.rule import RuleConfig
from app.services.document_parser import DocumentParser
from app.services.rule_engine import RuleEngine
from app.core.rules_config import get_rules_for_type

router = APIRouter(prefix="/api/documents", tags=["documents"])

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

from app.api.auth import get_current_user
from app.models.user import User

import uuid

MAX_FILE_SIZE = 10 * 1024 * 1024

@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    doc_type: str = Form("academic"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .docx ở phiên bản hiện tại")

    # Kiểm tra kích thước
    file_content = await file.read()
    if len(file_content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="Kích thước file vượt quá giới hạn 10MB")

    # Tạo tên an toàn
    safe_filename = f"{uuid.uuid4().hex}.docx"
    file_path = os.path.join(UPLOAD_DIR, safe_filename)
    
    # Lưu file
    with open(file_path, "wb") as buffer:
        buffer.write(file_content)
        
    try:
        # Lưu vào DB
        new_doc = Document(filename=safe_filename, status="processing", owner_id=current_user.id)
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)
        
        # Phân tích tài liệu
        parser = DocumentParser(file_path)
        doc_info = parser.extract_document_info()
        
        # Lấy rules theo doc_type và đánh giá
        if doc_type.startswith("custom_") or doc_type.startswith("rule_"):
            rule_id = int(doc_type.split("_")[1])
            rule_record = db.query(RuleConfig).filter(RuleConfig.id == rule_id).first()
            if not rule_record:
                raise HTTPException(status_code=404, detail="Không tìm thấy rule")
            rules = rule_record.rules_json
        else:
            rules = get_rules_for_type(doc_type)
            
        engine = RuleEngine(doc_info, rules)
        report = engine.evaluate()
        
        # Cập nhật trạng thái
        new_doc.status = "completed"
        db.commit()
        
        return {
            "document_id": new_doc.id,
            "filename": file.filename,
            "report": report
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

from fastapi.responses import FileResponse
from pydantic import BaseModel
import urllib.request

@router.get("/download/{document_id}")
def download_document(document_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu")
        
    file_path = os.path.join(UPLOAD_DIR, doc.filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File vật lý không tồn tại")
        
    return FileResponse(file_path)

class OnlyOfficeCallback(BaseModel):
    status: int
    url: str = None

@router.post("/callback/{document_id}")
async def onlyoffice_callback(document_id: int, payload: dict, db: Session = Depends(get_db)):
    # ONLYOFFICE gửi dữ liệu dạng application/json
    status = payload.get("status")
    
    # Status 2 nghĩa là tài liệu đã được chỉnh sửa và lưu
    if status == 2:
        download_url = payload.get("url")
        doc = db.query(Document).filter(Document.id == document_id).first()
        if doc and download_url:
            file_path = os.path.join(UPLOAD_DIR, doc.filename)
            # Tải file mới từ ONLYOFFICE ghi đè lên file cũ
            urllib.request.urlretrieve(download_url, file_path)
            
    return {"error": 0}

from docx import Document as DocxDocument
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from typing import List, Any

class FixItem(BaseModel):
    id: str
    type: str
    property: str
    value: Any
    para_index: int = None
    run_index: int = None

class FixRequest(BaseModel):
    fixes: List[FixItem]

@router.post("/fix/{document_id}")
async def fix_document(document_id: int, request: FixRequest, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    doc_record = db.query(Document).filter(Document.id == document_id).first()
    if not doc_record:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu")
    if doc_record.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Không có quyền chỉnh sửa tài liệu này")
        
    file_path = os.path.join(UPLOAD_DIR, doc_record.filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File vật lý không tồn tại")
        
    doc = DocxDocument(file_path)
    
    # Thực hiện sửa lỗi theo request
    for fix in request.fixes:
        if fix.type == "margin":
            for section in doc.sections:
                if fix.property == "top":
                    section.top_margin = Cm(fix.value)
                elif fix.property == "bottom":
                    section.bottom_margin = Cm(fix.value)
                elif fix.property == "left":
                    section.left_margin = Cm(fix.value)
                elif fix.property == "right":
                    section.right_margin = Cm(fix.value)
                    
        elif fix.type == "paragraph" and fix.para_index is not None:
            if fix.para_index < len(doc.paragraphs):
                para = doc.paragraphs[fix.para_index]
                if fix.property == "line_spacing":
                    para.paragraph_format.line_spacing = float(fix.value)
                elif fix.property == "spacing_before":
                    para.paragraph_format.space_before = Pt(float(fix.value))
                elif fix.property == "spacing_after":
                    para.paragraph_format.space_after = Pt(float(fix.value))
                elif fix.property == "alignment":
                    if fix.value == "left": para.alignment = 0
                    elif fix.value == "center": para.alignment = 1
                    elif fix.value == "right": para.alignment = 2
                    elif fix.value == "justified": para.alignment = 3
                    
        elif fix.type == "typography" and fix.para_index is not None and fix.run_index is not None:
            if fix.para_index < len(doc.paragraphs):
                para = doc.paragraphs[fix.para_index]
                if fix.run_index < len(para.runs):
                    run = para.runs[fix.run_index]
                    if fix.property == "font_name":
                        run.font.name = fix.value
                        # Đảm bảo font chữ hỗ trợ tốt tiếng Việt / Unicode
                        if run._element.rPr is not None and hasattr(run._element.rPr, 'rFonts'):
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), fix.value)
                    elif fix.property == "font_size":
                        run.font.size = Pt(fix.value)
                    elif fix.property == "bold":
                        run.font.bold = bool(fix.value)
                    elif fix.property == "italic":
                        run.font.italic = bool(fix.value)
                    elif fix.property == "uppercase":
                        run.font.all_caps = bool(fix.value)
                        
    # Lưu file đã sửa
    doc.save(file_path)
    
    return {"message": "Sửa lỗi thành công", "filename": doc_record.filename}
